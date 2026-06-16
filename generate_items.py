import io
import os
import re
import csv
import json
import time
from datetime import datetime
import pandas as pd
import openpyxl
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from mistralai.client import Mistral

from curriculum_analyzer import load_curriculum_framework, get_next_sequence_number

# ─── IMAGE EXTRACTION & GOOGLE DRIVE HELPERS ──────────────────────────────────
DRIVE_URL_PATTERN = re.compile(
    r'https?://(?:drive|docs)\.google\.com/[^\s\'"<>]+'
)

def download_google_drive_file(url):
    import requests
    # Extract file ID from URL
    file_id = None
    
    # Pattern 1: /file/d/FILE_ID/view
    m = re.search(r'/file/d/([a-zA-Z0-9_-]+)', url)
    if m:
        file_id = m.group(1)
    else:
        # Pattern 2: id=FILE_ID
        m = re.search(r'[?&]id=([a-zA-Z0-9_-]+)', url)
        if m:
            file_id = m.group(1)
            
    if not file_id:
        return None
        
    download_url = "https://docs.google.com/uc?export=download"
    session = requests.Session()
    
    try:
        response = session.get(download_url, params={'id': file_id}, stream=True, timeout=15)
        # Check for confirmation page for large files
        token = None
        for key, value in response.cookies.items():
            if key.startswith('download_warning'):
                token = value
                break
                
        if token:
            response = session.get(download_url, params={'id': file_id, 'confirm': token}, stream=True, timeout=15)
            
        if response.status_code == 200:
            return response.content
    except Exception as e:
        print(f"Error downloading Google Drive file {url}: {e}")
        
    return None

def clean_extracted_url(url):
    # Strip common punctuation that might be appended to the URL by regex
    while url and url[-1] in '.,;:?!':
        url = url[:-1]
    
    # Strip closing brackets/parentheses if they are not balanced in the URL
    while url and url.endswith(')') and url.count('(') < url.count(')'):
        url = url[:-1]
    while url and url.endswith(']') and url.count('[') < url.count(']'):
        url = url[:-1]
    while url and url.endswith('}') and url.count('{') < url.count('}'):
        url = url[:-1]
        
    return url

def extract_and_download_drive_images(text):
    if not text or not isinstance(text, str):
        return [], text
        
    urls = DRIVE_URL_PATTERN.findall(text)
    images = []
    cleaned_text = text
    
    for url in urls:
        cleaned_url = clean_extracted_url(url)
        img_bytes = download_google_drive_file(cleaned_url)
        if img_bytes:
            images.append(img_bytes)
            # Remove cleaned_url from the text and clean trailing/leading spaces or punctuation
            cleaned_text = cleaned_text.replace(cleaned_url, "").strip() if cleaned_url == url else cleaned_text.replace(cleaned_url, "").strip()
            
    # Clean up empty parentheses or brackets that might contain URLs, e.g. "()", "[]"
    cleaned_text = re.sub(r'\(\s*\)', '', cleaned_text)
    cleaned_text = re.sub(r'\[\s*\]', '', cleaned_text)
    cleaned_text = re.sub(r'\{\s*\}', '', cleaned_text)
    cleaned_text = cleaned_text.strip()
    
    return images, cleaned_text

def get_images_for_cell(sheet, row_idx, col_idx):
    """
    Returns a list of image bytes for images anchored in the cell at (row_idx, col_idx).
    row_idx and col_idx are 1-based (Excel style).
    """
    found_images = []
    if not hasattr(sheet, '_images'):
        return found_images
        
    for img in sheet._images:
        anchor = img.anchor
        img_row, img_col = None, None
        
        if hasattr(anchor, '_from'):
            img_row = anchor._from.row + 1
            img_col = anchor._from.col + 1
        elif isinstance(anchor, str):
            from openpyxl.utils import coordinate_to_tuple
            try:
                img_row, img_col = coordinate_to_tuple(anchor)
            except Exception:
                pass
                
        if img_row == row_idx and img_col == col_idx:
            try:
                data = img._data()
                found_images.append(data)
            except Exception as e:
                print(f"Error reading image data at row {row_idx}, col {col_idx}: {e}")
                
    return found_images

# ─── MISTRAL CLIENT SETUP ──────────────────────────────────────────────────────
MODEL = "mistral-large-latest"

def get_mistral_client(api_key: str):
    if api_key.lower() in ("mock", "test", ""):
        return None
    try:
        return Mistral(api_key=api_key)
    except Exception:
        return None

def get_ai_client(api_key: str, api_provider: str = "Mistral"):
    if api_key.lower() in ("mock", "test", ""):
        return None
    try:
        if api_provider.lower() == "mistral":
            return Mistral(api_key=api_key)
        elif api_provider.lower() == "groq":
            from groq import Groq
            return Groq(api_key=api_key)
    except Exception:
        return None

# ─── XML HELPERS FOR DOCX FORMATTING ──────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove existing shading if any
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_borders_all(cell, color_hex, sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_borders = tcPr.find(qn('w:tcBorders'))
    if existing_borders is not None:
        tcPr.remove(existing_borders)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), sz)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color_hex)
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing cell margins if any
    existing_margins = tcPr.find(qn('w:tcMar'))
    if existing_margins is not None:
        tcPr.remove(existing_margins)
    tcMar = OxmlElement('w:tcMar')
    for name, w in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(w))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_width_dxa(cell, w_dxa):
    cell.width = Pt(int(w_dxa) / 20)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(w_dxa))
    tcW.set(qn('w:type'), 'dxa')
    existing_tcW = tcPr.find(qn('w:tcW'))
    if existing_tcW is not None:
        tcPr.remove(existing_tcW)
    tcPr.append(tcW)

def create_styled_table(doc, num_rows, widths_dxa, border_color="cccccc"):
    num_cols = len(widths_dxa)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Table-level borders setup
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for b_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), border_color)
        tblBorders.append(b)
    tblPr.append(tblBorders)
    
    # Set total table width
    total_w = sum(widths_dxa)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_w))
    tblW.set(qn('w:type'), 'dxa')
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblPr.append(tblW)
    
    # Set column grid (tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_dxa:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    existing_tblGrid = table._tbl.find(qn('w:tblGrid'))
    if existing_tblGrid is not None:
        table._tbl.remove(existing_tblGrid)
    table._tbl.insert(1, tblGrid)
    
    # Format each cell
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        for idx, cell in enumerate(row.cells):
            set_cell_width_dxa(cell, widths_dxa[idx])
            set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
            set_cell_borders_all(cell, border_color)
            
    return table

# ─── PARAGRAPH HELPERS ────────────────────────────────────────────────────────
def add_section_header(doc, title, desc=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r1 = p.add_run(title)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0, 0, 0)
    
    if desc:
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)
        r2.font.bold = False
        r2.font.color.rgb = RGBColor(0, 0, 0)
    return p

def set_cell_text(cell, text, bold=False, italic=False, color_rgb="000000", alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if alignment is not None:
        p.alignment = alignment
        
    if not text:
        # Style empty run
        run = p.add_run("")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        return
        
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if alignment is not None:
                p.alignment = alignment
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = RGBColor.from_string(color_rgb)

# ─── JSON CLEANUP AND PARSE ───────────────────────────────────────────────────
def clean_and_parse_json(text: str):
    cleaned = text.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        cleaned = "\n".join(lines).strip()
    return json.loads(cleaned)

# ─── MISTRAL API AND RETRY SYSTEM ──────────────────────────────────────────────
def call_mistral(client, system_prompt, user_message, max_tokens=600, use_json=True, retries=3):
    if client is None:
        raise ValueError("Client not initialized (Mock Mode)")
        
    last_err = None
    client_type = client.__class__.__name__
    
    for attempt in range(retries):
        try:
            if client_type == "Mistral":
                response = client.chat.complete(
                    model="mistral-large-latest",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_message}
                    ],
                    temperature=0.1,
                    max_tokens=max_tokens,
                    response_format={"type": "json_object"} if use_json else None
                )
                content = response.choices[0].message.content
            elif client_type == "Groq":
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_message}
                    ],
                    temperature=0.1,
                    max_tokens=max_tokens,
                    response_format={"type": "json_object"} if use_json else None
                )
                content = response.choices[0].message.content
            else:
                raise ValueError(f"Unknown client type: {client_type}")
                
            if use_json:
                return clean_and_parse_json(content)
            else:
                return content.strip()
        except Exception as e:
            last_err = e
            time.sleep(1.5 * (attempt + 1))
            
    raise last_err

# ─── HIGH-FIDELITY MOCK FALLBACKS ──────────────────────────────────────────────
def generate_mock_mastery(row, inferred_type):
    # Try to determine from excel
    excel_m = str(row.get("Mastery_Level") or "").strip().upper()
    if excel_m in ("M1", "M2", "M3", "M4"):
        return {"mastery_level": excel_m, "reasoning": f"Validated mastery level {excel_m} from spreadsheet.", "operator_override": False}
        
    # Otherwise check type
    if inferred_type == "MCQ":
        m_level = "M2"
    elif inferred_type == "Extended Response":
        m_level = "M4"
    else:
        m_level = "M2"
        
    return {
        "mastery_level": m_level,
        "reasoning": f"Classified as {m_level} based on question structure (mock inference).",
        "operator_override": False
    }

def generate_mock_format_fields(row, item_type):
    stimulus = str(row.get("Stimulus_Text") or "")
    stem = str(row.get("Item_Stem") or "")
    combined = (stimulus + " " + stem).lower()

    # Detect image: keywords OR actual image filenames (e.g. img1_chp10.png)
    img_keyword = any(w in combined for w in ["image", "diagram", "figure", "picture", "photo"])
    img_filename = bool(re.search(r'\b[\w][\w.\-]*\.(?:png|jpg|jpeg|gif|svg|webp|bmp)\b', combined, re.IGNORECASE))
    has_image = "Yes" if img_keyword or img_filename else "No"

    # Detect table: matching table patterns (Column A/B) OR table keywords
    table_match = bool(re.search(r'(column\s+[ab]|match\s+the\s+following|col\s+[ab])', combined, re.IGNORECASE))
    table_keyword = any(w in combined for w in ["table", "data table", "grid"])
    has_table = "Yes" if table_match or table_keyword else "No"

    has_equation = "Yes" if any(w in combined for w in ["h2o", "co2", "formula", "equation", "°c", "m/s", "f=ma"]) else "No"
    eq_format = "LaTeX" if has_equation == "Yes" else "N/A"

    return {
        "has_image": has_image,
        "has_table": has_table,
        "has_equation": has_equation,
        "equation_format": eq_format
    }

def generate_mock_rationales(row, correct_letter, chapter_name="Science"):
    stem = str(row.get("Item_Stem") or "")
    opts = {
        "A": str(row.get("Option_A") or ""),
        "B": str(row.get("Option_B") or ""),
        "C": str(row.get("Option_C") or ""),
        "D": str(row.get("Option_D") or "")
    }
    # Provide a sequential mock error type so each distractor gets a distinct label
    error_types = ["Conceptual error", "Procedural error", "Comprehension error"]
    error_idx = 0

    rationales = {}
    for letter, text in opts.items():
        if letter == correct_letter:
            rationales[f"rationale_{letter.lower()}"] = (
                f"CORRECT ANSWER: {text} correctly identifies the scientific concept "
                f"being assessed. This option reflects an accurate understanding of the "
                f"process or classification described in the question stem."
            )
        else:
            error_type = error_types[error_idx % len(error_types)]
            error_idx += 1
            rationales[f"rationale_{letter.lower()}"] = (
                f"{error_type}: The learner selects '{text}' due to incomplete understanding "
                f"of the concept tested, applying a related but incorrect idea to the scenario."
            )

    return rationales

def generate_mock_explanation(row, item_type, correct_letter, rationales, chapter_name="Science"):
    stem = str(row.get("Item_Stem") or "")
    opts = {
        "A": str(row.get("Option_A") or ""),
        "B": str(row.get("Option_B") or ""),
        "C": str(row.get("Option_C") or ""),
        "D": str(row.get("Option_D") or "")
    }

    if item_type == "MCQ":
        correct_text = opts.get(correct_letter, "")
        # Required format: state the correct answer first (exact phrase)
        explanation = f"The correct answer is {correct_letter}. {correct_text}.\n\n"
        # Brief, accessible explanation of the concept behind the correct answer
        explanation += (
            "This option is correct because it accurately reflects the scientific concept "
            "being assessed. Understanding this idea helps you connect the topic to what "
            "you have learnt in class.\n\n"
        )
        # Required heading (exact phrase)
        explanation += "Why other options are incorrect:\n\n"
        # Per incorrect option in exact required format
        for letter in ["A", "B", "C", "D"]:
            if letter != correct_letter:
                text = opts.get(letter, "")
                rat = rationales.get(f"rationale_{letter.lower()}", "")
                # Strip error-type prefix for kid-friendly language, avoid double stops
                if ": " in rat and not rat.startswith("CORRECT"):
                    reason = rat.split(": ", 1)[1].rstrip(".")
                else:
                    reason = rat.rstrip(".")
                explanation += f"Option {letter} ({text}) is incorrect because {reason}.\n\n"
        return explanation.strip()
    else:
        # Constructed response: single cohesive paragraph applied to scenario
        ans = str(row.get("Correct_Answer") or "")
        scenario_ref = stem[:80].rstrip() if stem else "the scenario described in the question"
        return (
            f"{ans if ans else 'The correct response accurately addresses the scientific concept described in the question.'} "
            f"Applying this understanding to {scenario_ref} explains both the process involved "
            f"and the outcome described. Students who only partially answer this question may "
            f"identify one relevant factor but fail to connect it fully to the specific "
            f"situation described."
        )

def generate_mock_rubric(row, max_score, chapter_name="Science"):
    ans = str(row.get("Correct_Answer") or "")

    rows = []
    if max_score == 2:
        rows = [
            {
                "score": 2,
                "label": "Full marks",
                "criteria": (
                    "Student correctly identifies both key scientific concepts and demonstrates "
                    "understanding of their relationship, using appropriate scientific language."
                ),
                "sample": (
                    f"\"{ans if ans else 'The process happens because of the two factors described in the question, and without either one, the outcome changes.'}\""
                )
            },
            {
                "score": 1,
                "label": "Partial marks",
                "criteria": (
                    "Student correctly identifies one key concept OR identifies both concepts "
                    "but uses informal or imprecise language."
                ),
                "sample": "\"The process happens because of one of the main factors.\""
            },
            {
                "score": 0,
                "label": "Zero",
                "criteria": (
                    "Irrelevant answer, fundamental misconception, or purely observational "
                    "response with no scientific reasoning."
                ),
                "sample": "\"The process does not depend on any factors and happens by itself.\""
            }
        ]
    else:  # max_score == 3 or fallback
        rows = [
            {
                "score": 3,
                "label": "Full marks",
                "criteria": (
                    "Student correctly applies the relevant concept(s) to the scenario, "
                    "accurately predicts or explains the outcome, uses precise scientific "
                    "terminology, and demonstrates clear logical reasoning with no conceptual errors."
                ),
                "sample": (
                    f"\"{ans if ans else 'The scientific concept directly explains the outcome described in the scenario, as demonstrated by the clear process.'}\""
                )
            },
            {
                "score": 2,
                "label": "Partial",
                "criteria": (
                    "Correctly applies main concept(s) with minor omissions; correct outcome "
                    "stated with mostly complete reasoning; accurate in 2 of 3 required elements."
                ),
                "sample": "\"The process happens because of the main factor, but the connection is unclear.\""
            },
            {
                "score": 1,
                "label": "Minimal",
                "criteria": (
                    "Some relevant scientific understanding shown but concept applied incorrectly "
                    "or partially; OR correct outcome stated without reasoning."
                ),
                "sample": "\"It happens because of this factor.\""
            },
            {
                "score": 0,
                "label": "Zero",
                "criteria": "Irrelevant, fundamentally incorrect, or blank.",
                "sample": "\"Nothing will happen because the process does not depend on these things.\""
            }
        ]

    return {"rows": rows}

def get_matching_chapter_in_science_excel(selected_chap, excel_chapters):
    mapping = {
        "materials around us": "Materials Around Us",
        "living creatures": "Living Creatures: Exploring  their Characteristics",
        "diversity in the living world": "Diversity in the Living World",
        "exploring magnets": "Exploring Magnets",
        "mindful eating": "Mindful Eating: A Path  to a Healthy Body",
        "measurement of length and motion": "Measurement of Length and Motion",
        "temperature and its measurement": "Temperature and its Measurement",
        "a journey through states of water": "A Journey through States of Water",
        "methods of separation": "Method of Separation In Everyday Life",
        "nature's treasures": "Nature’s Treasures",
        "nature’s treasures": "Nature’s Treasures",
        "beyond earth": "Beyond Earth"
    }
    
    sel_norm = str(selected_chap).lower().strip().replace("’", "'").replace("`", "'")
    if sel_norm in mapping:
        mapped_name = mapping[sel_norm]
        for ec in excel_chapters:
            if ec.lower().strip().replace("’", "'").replace("`", "'") == mapped_name.lower().replace("’", "'"):
                return ec
    
    for ec in excel_chapters:
        ec_norm = ec.lower().strip().replace("’", "'").replace("`", "'")
        if sel_norm in ec_norm or ec_norm in sel_norm:
            return ec
            
    sel_words = set(sel_norm.split())
    for ec in excel_chapters:
        ec_words = set(ec.lower().strip().replace("’", "'").replace("`", "'").split())
        if sel_words.intersection(ec_words):
            return ec
            
    return selected_chap

def match_topic_heuristically(row_topic_or_stem, predefined_topics):
    if not predefined_topics:
        return None
    
    row_topic_norm = str(row_topic_or_stem).lower().strip()
    
    best_match = None
    best_score = -1
    
    for pt in predefined_topics:
        topic_name = pt['Topic']
        topic_norm = str(topic_name).lower().strip()
        
        if topic_norm == row_topic_norm:
            return pt
            
        if topic_norm in row_topic_norm or row_topic_norm in topic_norm:
            score = len(topic_norm)
            if score > best_score:
                best_score = score
                best_match = pt
                
    if best_match:
        return best_match
        
    for pt in predefined_topics:
        topic_name = pt['Topic']
        topic_words = set(re.findall(r'\w+', str(topic_name).lower()))
        input_words = set(re.findall(r'\w+', str(row_topic_or_stem).lower()))
        overlap = len(topic_words.intersection(input_words))
        if overlap > best_score:
            best_score = overlap
            best_match = pt
            
    return best_match or predefined_topics[0]

def classify_topic_with_ai(client, item_stem, stimulus_text, predefined_topics, row_topic=None):
    if not predefined_topics:
        return None
        
    if client is None:
        query = row_topic if row_topic else item_stem
        return match_topic_heuristically(query, predefined_topics)
        
    topics_list_str = ""
    for pt in predefined_topics:
        topics_list_str += f"- {pt['Topic']}\n"
        
    system_prompt = (
        "You are an educational curriculum mapping AI.\n"
        "Your task is to classify a science assessment item into exactly one topic from the predefined list of topics below.\n"
        "You must choose the topic that best matches the scientific concept tested in the question.\n\n"
        "PREDEFINED TOPICS:\n"
        f"{topics_list_str}\n"
        "INSTRUCTIONS:\n"
        "1. Choose the topic from the predefined list that is the best match for the question.\n"
        "2. Your output MUST be a JSON object containing exactly the key 'selected_topic' with the exact name of the topic from the list.\n"
        "3. Do not invent any new topics. Only select from the provided list."
    )
    
    user_message = (
        f"Item Stem: {item_stem}\n"
        f"Stimulus: {stimulus_text if stimulus_text else 'None'}\n"
    )
    if row_topic:
        user_message += f"Original Topic suggestion from Excel: {row_topic}\n"
        
    try:
        res = call_mistral(client, system_prompt, user_message, max_tokens=150, use_json=True)
        selected_name = res.get("selected_topic", "").strip()
        
        for pt in predefined_topics:
            if pt['Topic'].lower().strip() == selected_name.lower().strip():
                return pt
                
        for pt in predefined_topics:
            if selected_name.lower() in pt['Topic'].lower() or pt['Topic'].lower() in selected_name.lower():
                return pt
    except Exception as e:
        print(f"Topic classification AI call failed: {e}")
        
    query = row_topic if row_topic else item_stem
    return match_topic_heuristically(query, predefined_topics)

def get_framework_alignment_for_topic(topic_name, chapter_name, unique_combos):
    norm_topic = str(topic_name).lower().strip().replace("’", "'").replace("`", "'")
    norm_chap = str(chapter_name).lower().strip().replace("’", "'").replace("`", "'")
    
    for combo in unique_combos:
        c_topic = str(combo["topic_name"]).lower().strip().replace("’", "'").replace("`", "'")
        c_chap = str(combo["chapter_name"]).lower().strip().replace("’", "'").replace("`", "'")
        if c_topic == norm_topic and (norm_chap in c_chap or c_chap in norm_chap):
            return combo
            
    for combo in unique_combos:
        c_topic = str(combo["topic_name"]).lower().strip().replace("’", "'").replace("`", "'")
        if c_topic == norm_topic:
            return combo
            
    for combo in unique_combos:
        c_topic = str(combo["topic_name"]).lower().strip().replace("’", "'").replace("`", "'")
        if norm_topic in c_topic or c_topic in norm_topic:
            return combo
            
    for combo in unique_combos:
        c_chap = str(combo["chapter_name"]).lower().strip().replace("’", "'").replace("`", "'")
        if norm_chap in c_chap or c_chap in norm_chap:
            return combo
            
    return None

def compute_lo_id(combo, unique_combos):
    comp = combo.get("competency", "C-1.1")
    lo = combo.get("learning_outcome", "")
    
    comp_num = "1.1"
    m_comp = re.search(r"C-(\d+\.\d+)", comp)
    if m_comp:
        comp_num = m_comp.group(1)
    else:
        m_num = re.search(r"(\d+\.\d+)", comp)
        if m_num:
            comp_num = m_num.group(1)
            
    comp_los = []
    seen_lo = set()
    for entry in unique_combos:
        if entry["competency"] == comp:
            lo_clean = entry["learning_outcome"].strip()
            if lo_clean not in seen_lo:
                seen_lo.add(lo_clean)
                comp_los.append(lo_clean)
                
    try:
        curr_lo = lo.strip()
        match_lo_idx = 0
        for idx, entry_lo in enumerate(comp_los):
            if entry_lo.lower().strip() == curr_lo.lower().strip():
                match_lo_idx = idx
                break
        suffix_letter = chr(ord('a') + match_lo_idx)
    except Exception:
        suffix_letter = 'a'
        
    return f"LO-{comp_num}.{suffix_letter}"

# ─── CORE PIPELINE RUNNER ─────────────────────────────────────────────────────
def run_pipeline(excel_file, api_key: str, batch_meta: dict, progress_callback, api_provider: str = "Mistral"):
    """
    Ingests, validates, enriches and generates DOCX and CSV log.
    Returns: docx_bytes, csv_bytes, summary_dict
    """
    # 1. Load Excel
    try:
        wb = openpyxl.load_workbook(excel_file)
    except Exception as e:
        raise ValueError(f"Could not load Excel file: {e}")
        
    matching_sheet = None
    for name in wb.sheetnames:
        if name.strip().lower() == "questions":
            matching_sheet = name
            break
            
    if matching_sheet is None and len(wb.sheetnames) > 0:
        matching_sheet = wb.sheetnames[0]
        
    if matching_sheet is None:
        raise ValueError("Excel file is empty and has no sheets")
        
    sheet = wb[matching_sheet]
    
    # Load Science Class 6 mapping
    parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) if '__file__' in locals() else os.path.dirname(os.getcwd())
    science_path = os.path.join(parent_dir, "Science Class 6.xlsx")
    if not os.path.exists(science_path):
        science_path = os.path.join(os.path.dirname(os.getcwd()), "Science Class 6.xlsx")
    if not os.path.exists(science_path):
        science_path = "Science Class 6.xlsx"
        
    try:
        df_science = pd.read_excel(science_path)
    except Exception as e:
        raise ValueError(f"Could not load Science Class 6.xlsx: {e}")
        
    # Get the list of all unique chapters in the Science Class 6 sheet
    excel_chapters = df_science['Chapter'].dropna().unique().tolist()
    
    # Map batch_meta["chapter_name"] to the exact name in the excel
    matched_chapter = get_matching_chapter_in_science_excel(batch_meta["chapter_name"], excel_chapters)
    
    # Filter the Science Class 6 sheet to get predefined topics for this chapter
    df_chap = df_science[df_science['Chapter'] == matched_chapter]
    if df_chap.empty:
        df_chap = df_science[df_science['Chapter'].str.contains(matched_chapter[:10], case=False, na=False)]
        
    if df_chap.empty:
        raise ValueError(f"No topics found for chapter '{matched_chapter}' in Science Class 6.xlsx")
        
    predefined_topics = df_chap[['Topic ID', 'Topic']].drop_duplicates().to_dict('records')
    subject_id = str(df_chap.iloc[0]['Subject ID']).strip() if not df_chap.empty else "65ba59aef233a16d51f7163d"
    chapter_id = str(df_chap.iloc[0]['Chapter ID']).strip() if not df_chap.empty else ""
    
    # Load Curriculum Framework dynamically for mapping NCF CG, Competency, etc.
    framework_path = os.path.join(parent_dir, "Final_Class 6_Science_Mastery Framework_ (1).xlsx")
    if not os.path.exists(framework_path):
        framework_path = os.path.join(os.path.dirname(os.getcwd()), "Final_Class 6_Science_Mastery Framework_ (1).xlsx")
    if not os.path.exists(framework_path):
        framework_path = "Final_Class 6_Science_Mastery Framework_ (1).xlsx"
        
    unique_combos = load_curriculum_framework(framework_path)
    
    import difflib
    
    file_name_topic = None
    file_name_topic_id = None
    file_name_chapter_id = None
    best_ratio = 0
    
    if hasattr(excel_file, "name"):
        base_name = os.path.splitext(os.path.basename(excel_file.name))[0]
        base_name_clean = re.sub(r'[^a-zA-Z0-9\s]', ' ', base_name).strip().lower()
        
        # Match against all unique topics for 6th grade
        all_topics_df = df_science[['Topic ID', 'Topic', 'Chapter ID']].drop_duplicates()
        for _, r in all_topics_df.iterrows():
            t_name = str(r['Topic'])
            t_clean = re.sub(r'[^a-zA-Z0-9\s]', ' ', t_name).strip().lower()
            
            # Exact or substring match gets priority 1.0
            if t_clean in base_name_clean or base_name_clean in t_clean:
                ratio = 1.0
            else:
                ratio = difflib.SequenceMatcher(None, base_name_clean, t_clean).ratio()
                
            if ratio >= 0.4 and ratio > best_ratio:
                best_ratio = ratio
                file_name_topic = t_name
                file_name_topic_id = r['Topic ID']
                file_name_chapter_id = r['Chapter ID']
                
    if file_name_chapter_id:
        chapter_id = str(file_name_chapter_id).strip()

    # Track sequence number per topic ID
    current_seq = {}

    
    # Check headers
    expected_headers = [
        "Mastery_Level", "Stimulus_Text", "Item_Stem",
        "Option_A", "Option_B", "Option_C", "Option_D",
        "Correct_Option", "Correct_Answer"
    ]
    
    sheet_headers = [cell.value for cell in sheet[1]]
    # Handle minor trailing whitespaces or casing in header validation
    clean_headers = [str(h).strip() for h in sheet_headers if h is not None]
    
    # Check if all expected headers exist
    for eh in expected_headers:
        if eh not in clean_headers:
            raise ValueError(f"Required column '{eh}' not found in 'Questions' sheet headers. Found: {clean_headers}")
            
    # Map headers to indices
    header_map = {str(cell.value).strip(): col_idx for col_idx, cell in enumerate(sheet[1], start=1) if cell.value is not None}
    
    # Read rows
    raw_rows = []
    for r_idx in range(2, sheet.max_row + 1):
        row_dict = {}
        row_empty = True
        for h_name, col_idx in header_map.items():
            val = sheet.cell(row=r_idx, column=col_idx).value
            if val is not None:
                row_empty = False
            row_dict[h_name] = val
        if not row_empty:
            row_dict["_excel_row_num"] = r_idx
            
            # Extract and download images from Google Drive links & openpyxl embedded drawings
            row_images = []
            
            # Check Item_Stem
            stem_val = row_dict.get("Item_Stem")
            if stem_val and isinstance(stem_val, str):
                drive_imgs, cleaned_stem = extract_and_download_drive_images(stem_val)
                row_images.extend(drive_imgs)
                row_dict["Item_Stem"] = cleaned_stem
            
            if "Item_Stem" in header_map:
                embedded_stem_imgs = get_images_for_cell(sheet, r_idx, header_map["Item_Stem"])
                row_images.extend(embedded_stem_imgs)
                
            # Check Stimulus_Text
            stimulus_val = row_dict.get("Stimulus_Text")
            if stimulus_val and isinstance(stimulus_val, str):
                drive_imgs, cleaned_stimulus = extract_and_download_drive_images(stimulus_val)
                row_images.extend(drive_imgs)
                row_dict["Stimulus_Text"] = cleaned_stimulus
                
            if "Stimulus_Text" in header_map:
                embedded_stimulus_imgs = get_images_for_cell(sheet, r_idx, header_map["Stimulus_Text"])
                row_images.extend(embedded_stimulus_imgs)
                
            row_dict["_images"] = row_images
            raw_rows.append(row_dict)
            
    total_raw = len(raw_rows)
    
    # 2. Ingest and Validate
    valid_items = []
    skipped_items = 0
    override_count = 0
    
    audit_logs = []
    
    client = get_ai_client(api_key, api_provider)
    is_mock = client is None
    
    for idx, row in enumerate(raw_rows):
        excel_row_num = row["_excel_row_num"]
        stem = str(row.get("Item_Stem") or "").strip()
        
        # Validation A: Stem must not be blank
        if not stem:
            skipped_items += 1
            audit_logs.append({
                "Row_Number": excel_row_num,
                "Item_ID": "N/A",
                "Status": "ERROR",
                "Message": "Skipped: Item Stem is blank.",
                "Mastery_Level": "N/A",
                "Item_Type": "N/A",
                "Overrides": ""
            })
            continue
            
        opt_a = str(row.get("Option_A") or "").strip()
        opt_b = str(row.get("Option_B") or "").strip()
        opt_c = str(row.get("Option_C") or "").strip()
        opt_d = str(row.get("Option_D") or "").strip()
        correct_opt = str(row.get("Correct_Option") or "").strip().upper()
        correct_ans = str(row.get("Correct_Answer") or "").strip()
        mastery_excel = str(row.get("Mastery_Level") or "").strip().upper()
        
        # Validate mastery level syntax if provided
        if mastery_excel and mastery_excel not in ("M1", "M2", "M3", "M4"):
            skipped_items += 1
            audit_logs.append({
                "Row_Number": excel_row_num,
                "Item_ID": "N/A",
                "Status": "ERROR",
                "Message": f"Skipped: Invalid Mastery Level '{mastery_excel}' (must be M1, M2, M3, or M4).",
                "Mastery_Level": mastery_excel,
                "Item_Type": "N/A",
                "Overrides": ""
            })
            continue
            
        # Inferred Item Type logic (Rule 1.3)
        has_options = bool(opt_a or opt_b or opt_c or opt_d)
        is_mcq_opt = correct_opt in ("A", "B", "C", "D")
        
        inferred_item_type = None
        if has_options and is_mcq_opt:
            inferred_item_type = "MCQ"
        elif not has_options:
            if mastery_excel == "M4":
                inferred_item_type = "Extended Response"
            elif mastery_excel == "M1":
                # M1 must be MCQ (Rule 1.3 constraint error). We override and log a warning.
                inferred_item_type = "MCQ"
            else: # M2, M3 or empty
                inferred_item_type = "Short Answer"
        else:
            # Partially filled options or bad correct option
            inferred_item_type = "MCQ"
            
        # Validation B: MCQ options check
        if inferred_item_type == "MCQ":
            if not (opt_a and opt_b and opt_c and opt_d):
                skipped_items += 1
                audit_logs.append({
                    "Row_Number": excel_row_num,
                    "Item_ID": "N/A",
                    "Status": "ERROR",
                    "Message": "Skipped: MCQ item is missing one or more options.",
                    "Mastery_Level": mastery_excel or "N/A",
                    "Item_Type": "MCQ",
                    "Overrides": ""
                })
                continue
            if not is_mcq_opt:
                skipped_items += 1
                audit_logs.append({
                    "Row_Number": excel_row_num,
                    "Item_ID": "N/A",
                    "Status": "ERROR",
                    "Message": f"Skipped: MCQ item has invalid Correct Option '{correct_opt}' (must be A, B, C, or D).",
                    "Mastery_Level": mastery_excel or "N/A",
                    "Item_Type": "MCQ",
                    "Overrides": ""
                })
                continue
                
        # All basic load validations passed!
        valid_items.append({
            "row": row,
            "inferred_item_type": inferred_item_type,
            "mastery_excel": mastery_excel
        })
        
    # Process valid items
    total_valid = len(valid_items)
    processed_items = []
    
    # Sequence numbers
    seq_num = batch_meta["start_seq"]
    
    for i, item in enumerate(valid_items):
        row = item["row"]
        excel_row_num = row["_excel_row_num"]
        inferred_item_type = item["inferred_item_type"]
        mastery_excel = item["mastery_excel"]
        topic = str(row.get("Topic") or "").strip()
        stimulus = str(row.get("Stimulus_Text") or "").strip()
        stem = str(row.get("Item_Stem") or "").strip()
        opt_a = str(row.get("Option_A") or "").strip()
        opt_b = str(row.get("Option_B") or "").strip()
        opt_c = str(row.get("Option_C") or "").strip()
        opt_d = str(row.get("Option_D") or "").strip()
        correct_opt = str(row.get("Correct_Option") or "").strip().upper()
        correct_ans = str(row.get("Correct_Answer") or "").strip()
        
        # ─── TOPIC CLASSIFICATION & ALIGNMENT ─────────────────────────────────
        if file_name_topic:
            item_topic = file_name_topic
            item_topic_id = file_name_topic_id
        else:
            pt = classify_topic_with_ai(client, stem, stimulus, predefined_topics, row_topic=topic)
            item_topic = pt['Topic']
            item_topic_id = pt['Topic ID']
        
        combo = get_framework_alignment_for_topic(item_topic, matched_chapter, unique_combos)
        if combo:
            item_ncf_cg = combo['ncf_cg']
            item_competency = combo['competency']
            item_learning_outcome = combo['learning_outcome']
            item_lo_id = compute_lo_id(combo, unique_combos)
        else:
            # Fallback to the first entry in unique_combos matching this chapter
            chap_fallback = None
            if unique_combos:
                norm_matched_chap = re.sub(r"[\s\n\r\t]+", " ", matched_chapter.lower().strip())
                for c in unique_combos:
                    norm_c_chap = re.sub(r"[\s\n\r\t]+", " ", c["chapter_name"].lower().strip())
                    if norm_matched_chap in norm_c_chap or norm_c_chap in norm_matched_chap:
                        chap_fallback = c
                        break
            if chap_fallback:
                item_ncf_cg = chap_fallback['ncf_cg']
                item_competency = chap_fallback['competency']
                item_learning_outcome = chap_fallback['learning_outcome']
                item_lo_id = compute_lo_id(chap_fallback, unique_combos)
            else:
                item_ncf_cg = "CG-1"
                item_competency = "C-1.1"
                item_learning_outcome = "Identifies and explains different properties of materials and relates them to their uses."
                item_lo_id = "LO-1.1.a"

        # Sequence number starts from start_seq and increments sequentially for each item in this run
        item_seq = i + batch_meta.get("start_seq", 1)

        # Temp ID for display (starts sequentially, no more -XXX)
        temp_id = f"G6-{chapter_id}-{item_topic_id}-{mastery_excel or 'MX'}-{item_seq:03d}"
        
        # Log update to Streamlit UI
        if progress_callback:
            progress_callback(i + 1, total_valid, temp_id, mastery_excel or "MX", "Classifying and validating cognitive demand...")
            
        # ─── CALL 1: MASTERY LEVEL CLASSIFICATION ──────────────────────────────
        confirmed_mastery = "M2" # Fallback default
        override_msg = ""
        
        if mastery_excel in ("M1", "M2", "M3", "M4"):
            confirmed_mastery = mastery_excel
        else:
            if is_mock:
                c1_res = generate_mock_mastery(row, inferred_item_type)
            else:
                try:
                    system_prompt_c1 = (
                        "You are an expert assessment designer for the NCF 2023 Grade 6 Science curriculum in India.\n"
                        "You classify assessment items into one of four Mastery Levels based strictly on the cognitive\n"
                        "demand required to answer correctly — not on topic difficulty.\n\n"
                        "MASTERY LEVEL DEFINITIONS:\n\n"
                        "M1 (Remembering, DoK 1):\n"
                        "  The learner identifies, names, or recognises from a given set. No explanation required.\n"
                        "  Direct recall or recognition. ALWAYS MCQ or matching format.\n\n"
                        "M2 (Understanding, DoK 2):\n"
                        "  The learner explains, classifies, or describes using scientific reasoning in a familiar,\n"
                        "  structured context. Requires understanding, not just recall. MCQ with justification or Short Answer.\n\n"
                        "M3 (Applying, DoK 3):\n"
                        "  The learner solves an unfamiliar scenario, applies knowledge to a novel situation, or reasons\n"
                        "  through a non-obvious case. The stimulus introduces context not directly taught.\n"
                        "  Can be scenario-based MCQ or Short Answer / Extended Response.\n\n"
                        "M4 (Evaluating, DoK 4 — Extended Thinking):\n"
                        "  The learner evaluates a flawed claim, incorrect reasoning, or competing explanation in the stimulus,\n"
                        "  identifies the specific error in REASONING (not just facts), and constructs a justified correction.\n"
                        "  ALWAYS Extended Response. The flaw must be in the reasoning, not just the outcome.\n\n"
                        "CLASSIFICATION RULES:\n"
                        "1. Base classification on the cognitive demand of STEM + STIMULUS together, not topic difficulty.\n"
                        "2. An item answerable purely from memory → M1 or M2, never M3 or M4.\n"
                        "3. M3 requires an unfamiliar scenario; applying knowledge in a familiar context → M2.\n"
                        "4. M4 REQUIRES a flawed reasoning claim in the stimulus. No flaw → cannot be M4.\n"
                        "5. If operator supplied a Mastery_Level, validate it. If incorrect, override and explain.\n\n"
                        "OUTPUT: JSON only, no other text:\n"
                        '{"mastery_level": "M1", "reasoning": "One sentence justifying the classification.", "operator_override": false}'
                    )
                    
                    correct_ans_sum = row["Correct_Option"] if inferred_item_type == "MCQ" else row["Correct_Answer"]
                    user_msg_c1 = (
                        f"Item Type (inferred): {inferred_item_type}\n"
                        f"Operator-supplied Mastery Level: {mastery_excel if mastery_excel else 'Not provided — please classify.'}\n"
                        f"Stimulus: {stimulus if stimulus else 'None'}\n"
                        f"Item Stem: {stem}\n"
                        f"Correct Answer / Option: {correct_ans_sum}"
                    )
                    
                    c1_res = call_mistral(client, system_prompt_c1, user_msg_c1, max_tokens=150, use_json=True)
                except Exception as exc:
                    fallback_level = mastery_excel if mastery_excel in ("M1", "M2", "M3", "M4") else "M2"
                    c1_res = {"mastery_level": fallback_level, "reasoning": f"Fallback to {fallback_level} due to API error: {exc}", "operator_override": False}
                    
            confirmed_mastery = c1_res.get("mastery_level", "M2").upper()
            if confirmed_mastery not in ("M1", "M2", "M3", "M4"):
                confirmed_mastery = "M2"
                
        if mastery_excel and confirmed_mastery != mastery_excel:
            override_msg = f"Override: Operator suggested {mastery_excel}, AI classified as {confirmed_mastery}."
            override_count += 1
            
        # ─── CONSTRAINT ENFORCEMENT & OVERRIDES (Rule 3.12) ───────────────────
        item_type = inferred_item_type
        
        if confirmed_mastery == "M4" and item_type != "Extended Response":
            item_type = "Extended Response"
            msg = "M4 items must be Extended Response. Inferred item type overridden to Extended Response."
            override_msg = (override_msg + " " + msg).strip()
            override_count += 1
            
        elif confirmed_mastery == "M1" and item_type != "MCQ":
            item_type = "MCQ"
            msg = "M1 items must be MCQ. Inferred item type overridden to MCQ."
            override_msg = (override_msg + " " + msg).strip()
            override_count += 1
            
        elif confirmed_mastery == "M2" and item_type == "Extended Response":
            item_type = "Short Answer"
            msg = "M2 items can only be MCQ or Short Answer. Inferred item type overridden to Short Answer."
            override_msg = (override_msg + " " + msg).strip()
            override_count += 1
            
        # Warn for missing stimulus on M3/M4 items
        if confirmed_mastery in ("M3", "M4") and not stimulus:
            msg = f"WARNING: {confirmed_mastery} item is missing a stimulus."
            override_msg = (override_msg + " " + msg).strip()
            
        # Finalized Item ID (using the sequence number resolved earlier)
        item_id = f"G6-{chapter_id}-{item_topic_id}-{confirmed_mastery}-{item_seq:03d}"
        
        # ─── CALL 2: FORMAT FIELD INFERENCE ──────────────────────────────────
        if progress_callback:
            progress_callback(i + 1, total_valid, item_id, confirmed_mastery, "Inferring rendering properties...")

        if is_mock:
            c2_res = generate_mock_format_fields(row, item_type)
        else:
            try:
                system_prompt_c2 = (
                    "You are a technical data-entry assistant for an educational assessment platform.\n"
                    "Analyse the item content and determine its rendering properties.\n\n"
                    "Has_Image:\n"
                    "  YES if stimulus or stem references an image, photograph, diagram, figure, picture,\n"
                    "  OR contains any image filename (e.g., 'img1_chp10.png', 'figure_2.jpg'),\n"
                    "  OR contains phrases like 'the diagram below', 'the figure above', 'look at the image'.\n"
                    "  NO otherwise.\n\n"
                    "Has_Table:\n"
                    "  YES if stimulus or stem contains a data table, matching table (Column A / Column B),\n"
                    "  'match the following', grid, or any explicit reference to tabular data.\n"
                    "  NO otherwise.\n\n"
                    "Has_Equation:\n"
                    "  YES if the item contains mathematical symbols, chemical formulae, physical quantities\n"
                    "  in symbolic form (e.g., H₂O, CO₂, 37°C, m/s², F=ma), or scientific notation.\n"
                    "  NO for plain-text descriptions.\n\n"
                    "Equation_Format:\n"
                    "  If Has_Equation is YES → 'LaTeX' if cleanly LaTeX-expressible; 'Image' if a visual\n"
                    "  is needed. 'N/A' if Has_Equation is NO.\n\n"
                    "OUTPUT: JSON only, no other text:\n"
                    '{"has_image": "Yes", "has_table": "No", "has_equation": "No", "equation_format": "N/A"}'
                )
                user_msg_c2 = (
                    f"Stimulus: {stimulus if stimulus else 'None'}\n"
                    f"Item Stem: {stem}\n"
                    f"Option A: {row.get('Option_A') if item_type == 'MCQ' else 'N/A'}\n"
                    f"Option B: {row.get('Option_B') if item_type == 'MCQ' else 'N/A'}\n"
                    f"Option C: {row.get('Option_C') if item_type == 'MCQ' else 'N/A'}\n"
                    f"Option D: {row.get('Option_D') if item_type == 'MCQ' else 'N/A'}"
                )
                c2_res = call_mistral(client, system_prompt_c2, user_msg_c2, max_tokens=120, use_json=True)
            except Exception as exc:
                c2_res = {"has_image": "No", "has_table": "No", "has_equation": "No", "equation_format": "N/A"}

        has_image = str(c2_res.get("has_image", "No")).strip().capitalize()
        if row.get("_images"):
            has_image = "Yes"
        has_table = str(c2_res.get("has_table", "No")).strip().capitalize()
        has_equation = str(c2_res.get("has_equation", "No")).strip().capitalize()

        # ─── POST-PROCESS: Detect image filenames embedded in stem/stimulus ───
        # e.g. "img1_chp10.png", "figure_3.jpg" → extract, mark Has_Image=Yes,
        # and remove the raw filename from the stem (per delivery checklist).
        IMG_FILENAME_RE = re.compile(
            r'\b([\w][\w.\-]*\.(?:png|jpg|jpeg|gif|svg|webp|bmp))\b',
            re.IGNORECASE
        )
        detected_image_filename = None
        all_img_filenames = IMG_FILENAME_RE.findall(stem) + IMG_FILENAME_RE.findall(stimulus)
        if all_img_filenames:
            has_image = "Yes"
            detected_image_filename = all_img_filenames[0]
            # Remove bare filenames from stem (leave surrounding context intact)
            for fn in IMG_FILENAME_RE.findall(stem):
                stem = re.sub(
                    r'[\(\[\s]*' + re.escape(fn) + r'[\)\]\s]*',
                    ' ', stem
                ).strip()
            stem = re.sub(r'\s{2,}', ' ', stem).strip()

        # ─── POST-PROCESS: Detect matching-table patterns ─────────────────────
        TABLE_DETECT_RE = re.compile(
            r'(column\s+[ab]|match\s+the\s+following|col\s+[ab]|column\s+[12])',
            re.IGNORECASE
        )
        if TABLE_DETECT_RE.search(stem + " " + (stimulus or "")):
            has_table = "Yes"

        # ─── Resolve Image_File_Name ──────────────────────────────────────────
        if detected_image_filename:
            image_file_name = detected_image_filename
        elif has_image == "Yes":
            image_file_name = f"{item_id}-img01.png"
        else:
            image_file_name = ""

        # Clean up eq_format to be LaTeX or Image if Has_Equation is Yes, else blank
        eq_format_raw = str(c2_res.get("equation_format", "")).strip()
        if has_equation == "Yes":
            if "latex" in eq_format_raw.lower():
                eq_format = "LaTeX"
            elif "image" in eq_format_raw.lower():
                eq_format = "Image"
            else:
                eq_format = "LaTeX"  # default fallback
        else:
            eq_format = ""
        
        # ─── CALL 3: DISTRACTOR RATIONALE GENERATION (MCQ Only) ───────────────
        distractor_rationales = {}
        if item_type == "MCQ":
            if progress_callback:
                progress_callback(i + 1, total_valid, item_id, confirmed_mastery, "Generating distractor rationales...")

            if is_mock:
                distractor_rationales = generate_mock_rationales(row, correct_opt, chapter_name=matched_chapter)
            else:
                try:
                    system_prompt_c3 = (
                        "You are an expert assessment designer for Grade 6 Science (NCF 2023, India).\n"
                        "Your task is to write DIAGNOSTIC OPTION RATIONALES for an MCQ.\n\n"
                        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                        "DIAGNOSTIC OPTION RATIONALE GUIDELINES\n"
                        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                        "For every MCQ, provide a specific, diagnostic rationale for ALL FOUR options.\n"
                        "Base your analysis on the exact scientific concepts tested and how a\n"
                        "6th-grade learner thinks about this specific topic.\n\n"
                        "1. CORRECT OPTION:\n"
                        "   Always begin with exactly: CORRECT ANSWER:\n"
                        "   Provide a clear, direct scientific explanation of why this option is correct.\n"
                        "   Be specific to the concept in this item — do NOT give a generic explanation.\n\n"
                        "2. INCORRECT OPTIONS (Distractors):\n"
                        "   Identify the exact cognitive misstep the learner made.\n"
                        "   Start with ONE of these three error types (exactly as written), followed by\n"
                        "   a colon, then explain the specific flawed thinking:\n\n"
                        "   Conceptual error: Errors due to misconceptions or factual errors. Stems from\n"
                        "     partial understanding or confusion between related concepts.\n"
                        "     Example: 'Conceptual error: Learner places leaves before germination,\n"
                        "     confusing when a seed first sprouts with later leaf growth.'\n\n"
                        "   Procedural error: Errors from incorrect application of a process or\n"
                        "     step-by-step observation. Used mainly for application/scenario questions.\n"
                        "     Example: 'Procedural error: Learner reverses the germination steps,\n"
                        "     placing root emergence after shoot appearance.'\n\n"
                        "   Comprehension error: Errors from misreading or misinterpreting the question\n"
                        "     language or a specific term/phrase in the stem.\n"
                        "     Example: 'Comprehension error: Learner reads \"next stage\" as referring to\n"
                        "     the mature plant rather than the immediate developmental step.'\n\n"
                        "CRITICAL RULES — VIOLATIONS WILL CAUSE REJECTION:\n"
                        "  ✗ NEVER write 'confused this concept with another related scientific idea'\n"
                        "  ✗ NEVER reference the chapter title (e.g., 'Living Creatures: Exploring their\n"
                        "    Characteristics') — reference the SPECIFIC concept being assessed instead\n"
                        "    (e.g., 'germination conditions', 'life cycle of a bean plant',\n"
                        "    'flower-to-pod/fruit formation')\n"
                        "  ✗ NEVER write generic rationales that could apply to any question\n"
                        "  ✓ ALWAYS name the exact misconception specific to this item\n"
                        "  ✓ Maximum 1-2 sentences per rationale\n"
                        "  ✓ Written for the item author, not the student\n\n"
                        "OUTPUT: JSON only, no other text:\n"
                        '{\n'
                        '  "rationale_a": "...",\n'
                        '  "rationale_b": "...",\n'
                        '  "rationale_c": "...",\n'
                        '  "rationale_d": "..."\n'
                        '}'
                    )

                    user_msg_c3 = (
                        f"Topic: {topic}\n"
                        f"Stimulus: {stimulus if stimulus else 'None'}\n"
                        f"Item Stem: {stem}\n"
                        f"Option A: {row.get('Option_A')} (Correct: {'Yes' if correct_opt == 'A' else 'No'})\n"
                        f"Option B: {row.get('Option_B')} (Correct: {'Yes' if correct_opt == 'B' else 'No'})\n"
                        f"Option C: {row.get('Option_C')} (Correct: {'Yes' if correct_opt == 'C' else 'No'})\n"
                        f"Option D: {row.get('Option_D')} (Correct: {'Yes' if correct_opt == 'D' else 'No'})\n\n"
                        "Write rationales that are SPECIFIC to this exact question. "
                        "For each wrong option, name the exact misconception a Grade 6 learner "
                        "would hold that makes that option attractive. Do NOT use generic phrases."
                    )

                    distractor_rationales = call_mistral(client, system_prompt_c3, user_msg_c3, max_tokens=800, use_json=True)
                except Exception:
                    distractor_rationales = generate_mock_rationales(row, correct_opt, chapter_name=matched_chapter)

        # ─── CALL 4: ANSWER EXPLANATION GENERATION ────────────────────────────
        if progress_callback:
            progress_callback(i + 1, total_valid, item_id, confirmed_mastery, "Generating answer explanation...")

        if is_mock:
            explanation_text = generate_mock_explanation(row, item_type, correct_opt, distractor_rationales, chapter_name=matched_chapter)
        else:
            try:
                system_prompt_c4 = (
                    "You are a Grade 6 Science teacher in India writing student-facing feedback\n"
                    "for an educational platform. Write in plain, simple, kid-friendly British English\n"
                    "suitable for 11-to-12-year-olds.\n\n"
                    "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                    "STUDENT-FACING ANSWER EXPLANATION GUIDELINES\n"
                    "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                    "IF THE QUESTION IS AN MCQ — use this EXACT structure (no deviations):\n"
                    "  Step 1 — State correct answer (exact phrase required):\n"
                    "    'The correct answer is [Option Letter]. [Full Option Text].'\n"
                    "  Step 2 — Explain the correct answer:\n"
                    "    Brief paragraph (2-3 sentences) explaining the scientific concept.\n"
                    "    Apply the concept DIRECTLY to the specific scenario/question asked.\n"
                    "    Do NOT give a generic definition. Connect it to what the question described.\n"
                    "  Step 3 — Distractor heading (exact phrase required):\n"
                    "    'Why other options are incorrect:'\n"
                    "  Step 4 — Per incorrect option (exact format required):\n"
                    "    'Option [Letter] ([Full Option Text]) is incorrect because [simple, kid-friendly\n"
                    "    explanation of the specific misconception].'\n"
                    "    Repeat for each incorrect option.\n\n"
                    "IF THE QUESTION IS CONSTRUCTED RESPONSE (Short Answer / ECR) — use this structure:\n"
                    "  Write a SINGLE cohesive paragraph.\n"
                    "  Do NOT just state the generic scientific definition.\n"
                    "  Apply the scientific explanation directly to the specific characters, objects,\n"
                    "  or scenario mentioned in the question (e.g., if the question mentions Riya's\n"
                    "  diagram, refer to Riya's diagram specifically).\n\n"
                    "LANGUAGE RULES:\n"
                    "  - Plain British English. Short sentences. Age 11-12.\n"
                    "  - Scientifically accurate, NCF 2023 Grade 6 aligned.\n"
                    "  - Everyday analogies where helpful.\n"
                    "  - NEVER end a sentence with double full stops (..).\n"
                    "  - NEVER reference the chapter title generically (e.g., 'Living Creatures:\n"
                    "    Exploring their Characteristics'). Reference the specific concept instead.\n\n"
                    "OUTPUT: Plain text only. No JSON. No markdown headers. Just the explanation text."
                )

                correct_ans_or_opt_text = (
                    row["Correct_Option"] + ". " + str(row.get(f"Option_{correct_opt}"))
                    if item_type == "MCQ" else row["Correct_Answer"]
                )

                user_msg_c4 = (
                    f"Item Type: {item_type}\n"
                    f"Topic: {topic}\n"
                    f"Stimulus: {stimulus if stimulus else 'None'}\n"
                    f"Item Stem: {stem}\n"
                    f"Correct Answer: {correct_ans_or_opt_text}\n"
                    + (
                        f"Option A: {row.get('Option_A')}\n"
                        f"Option B: {row.get('Option_B')}\n"
                        f"Option C: {row.get('Option_C')}\n"
                        f"Option D: {row.get('Option_D')}\n"
                        if item_type == "MCQ" else ""
                    )
                    + (
                        f"Distractor rationales (use these to write the 'Why other options are "
                        f"incorrect' section): {json.dumps(distractor_rationales)}"
                        if item_type == "MCQ" else ""
                    )
                )

                explanation_text = call_mistral(client, system_prompt_c4, user_msg_c4, max_tokens=700, use_json=False)
            except Exception:
                explanation_text = generate_mock_explanation(row, item_type, correct_opt, distractor_rationales, chapter_name=matched_chapter)
                
        # ─── DETERMINISTIC RULES DERIVATIONS ──────────────────────────────────
        blooms = {
            "M1": "Remembering",
            "M2": "Understanding",
            "M3": "Applying",
            "M4": "Evaluating"
        }.get(confirmed_mastery, "Understanding")
        
        dok = {
            "M1": "1",
            "M2": "2",
            "M3": "3",
            "M4": "4"  # M4 = Evaluating + Extended Thinking
        }.get(confirmed_mastery, "2")
        
        # Est Time
        if item_type == "MCQ":
            est_time = "1 min" if confirmed_mastery == "M1" else "2 mins"
        elif item_type == "Short Answer":
            est_time = "2 mins"
        else: # Extended Response
            est_time = "5 mins"
            
        # Max Score
        if item_type == "MCQ":
            max_score = 1
        elif confirmed_mastery == "M2":
            max_score = 2
        else: # M3 and M4
            max_score = 3
            
        # Scoring Type
        scoring_type = "Dichotomous (0/1)" if item_type == "MCQ" else f"Polytomous (0-{max_score})"
        
        # Scale
        if item_type == "MCQ":
            scale = ""
        else:
            scale = ", ".join(str(s) for s in range(max_score + 1))
            
        # Prerequisite Level
        prereq = {
            "M1": "",
            "M2": "M1",
            "M3": "M2",
            "M4": "M3"
        }.get(confirmed_mastery, "")
        
        # Options Count
        num_options = "4" if item_type == "MCQ" else ""
        
        # Feedback type
        feedback_type = "Immediate" if item_type == "MCQ" else "Delayed"
        
        # ─── CALL 5: MARKING RUBRIC GENERATION (Non-MCQ Only) ─────────────────
        rubric_data = {"rows": []}
        if item_type != "MCQ":
            if progress_callback:
                progress_callback(i + 1, total_valid, item_id, confirmed_mastery, "Generating marking rubric...")

            if is_mock:
                rubric_data = generate_mock_rubric(row, max_score, chapter_name=matched_chapter)
            else:
                try:
                    system_prompt_c5 = (
                        "You are a marking rubric designer for Grade 6 Science assessments aligned to NCF 2023 India.\n"
                        "Your rubrics follow the partial-credit scoring model: binary (full or zero) marking is\n"
                        "NOT acceptable. Every rubric must have a score tier for each point from 0 to the max score.\n\n"
                        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                        "CRITICAL FORMATTING RULES — STRICTLY ENFORCED:\n"
                        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                        "1. SAMPLE RESPONSE TONE: Write direct, confident answers that sound like a 6th-grade\n"
                        "   student writing on a formal science test. DO NOT write conversational, childish,\n"
                        "   or self-deprecating dialogue. BANNED PHRASES: \"I think...\", \"I don't know...\",\n"
                        "   \"I'm not sure...\", \"It magically happens.\"\n\n"
                        "2. INVERTED COMMAS: Every Sample Response MUST begin with an opening double\n"
                        "   quotation mark (\") and end with a closing double quotation mark (\").\n\n"
                        "3. SAMPLE RESPONSE LEVELS:\n"
                        "   - Full Marks: Direct, scientifically accurate sentence perfectly hitting the rubric\n"
                        "     criteria using correct terminology. (e.g. \"The bean seed undergoes germination.\")\n"
                        "   - Partial Marks: Direct, confident-sounding answer missing key info or using basic\n"
                        "     layman terms. (e.g. \"The seed opens up and grows a new plant when it gets wet.\")\n"
                        "   - Zero Marks: Direct, confident-sounding answer with a complete factual error or\n"
                        "     misconception. Do NOT make them sound confused. (e.g. \"The seed melts into the dirt.\")\n\n"
                        "4. ITEM-SPECIFIC CRITERIA: Every criteria row must mention the exact scoring\n"
                        "   indicators for THIS question — not generic descriptions.\n"
                        "   ✗ NEVER write: 'correct terminology for Living Creatures: Exploring their\n"
                        "     Characteristics' or 'appropriate properties of Living Creatures'\n"
                        "   ✓ DO write: 'identifies the missing fruit/pod-with-seeds stage',\n"
                        "     'explains role of water and air in seed germination',\n"
                        "     'describes flower-to-pod formation in the bean plant life cycle'\n\n"
                        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                        "RUBRIC TIER DEFINITIONS (apply based on max score):\n"
                        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                        "FOR MAX SCORE = 2 (Short Answer, M2) — 3 rows:\n"
                        "  2 marks (Full marks): Student correctly identifies BOTH required scientific\n"
                        "    concepts using appropriate terminology AND demonstrates understanding of\n"
                        "    their relationship or application to the scenario.\n"
                        "  1 mark (Partial marks): Student correctly identifies ONE concept; OR identifies\n"
                        "    both but with informal/incorrect terminology; OR shows partial understanding.\n"
                        "  0 marks (Zero): Irrelevant, fundamental misconception, or purely observational\n"
                        "    response with no scientific reasoning.\n\n"
                        "FOR MAX SCORE = 3 (Short Answer or Extended Response, M3 and M4) — 4 rows:\n"
                        "  3 marks (Full marks): Student correctly applies all concept(s) to the scenario,\n"
                        "    accurately explains outcome, uses precise terminology, clear logical reasoning.\n"
                        "  2 marks (Partial): Correct outcome with mostly complete reasoning; accurate\n"
                        "    in 2 of 3 required elements; minor omissions only.\n"
                        "  1 mark (Minimal): Some relevant understanding shown but applied incorrectly or\n"
                        "    partially; OR correct outcome stated without reasoning.\n"
                        "  0 marks (Zero): Irrelevant, fundamentally incorrect, or blank.\n\n"
                        "OUTPUT: JSON only, no other text:\n"
                        "{\n"
                        '  "rows": [\n'
                        '    {"score": 3, "label": "Full marks", "criteria": "...", "sample": "\\\"Student response here...\\\""},\n'
                        '    {"score": 2, "label": "Partial", "criteria": "...", "sample": "\\\"Student response here...\\\""},\n'
                        '    {"score": 1, "label": "Minimal", "criteria": "...", "sample": "\\\"Student response here...\\\""},\n'
                        '    {"score": 0, "label": "Zero", "criteria": "...", "sample": "\\\"Student response here...\\\"" }\n'
                        '  ]\n'
                        "}\n"
                        "Score values: 0 through max_score in descending order.\n"
                        'Labels: "Full marks", "Partial" (or "Partial marks"), "Minimal", "Zero".'
                    )

                    user_msg_c5 = (
                        f"Item Type: {item_type}\n"
                        f"Mastery Level: {confirmed_mastery}\n"
                        f"Max Score: {max_score}\n"
                        f"Topic: {topic}\n"
                        f"Item Stem: {stem}\n"
                        f"Stimulus: {stimulus if stimulus else 'None'}\n"
                        f"Model Answer / Correct Response: {row.get('Correct_Answer')}\n\n"
                        "REMINDER: Sample responses must be in double quotation marks and written as "
                        "direct, confident answers. Criteria must reference the exact concepts in "
                        "THIS question — never use generic phrases or chapter-title references."
                    )

                    rubric_data = call_mistral(client, system_prompt_c5, user_msg_c5, max_tokens=1200, use_json=True)
                except Exception:
                    rubric_data = generate_mock_rubric(row, max_score, chapter_name=matched_chapter)
                    
            # Rubric Quality Check (Rule 7.6)
            rows = rubric_data.get("rows", [])
            valid_rows = []
            expected_scores = list(range(max_score, -1, -1))
            
            # Rebuild clean rows based on expected scores to guarantee correct structure
            for exp_score in expected_scores:
                # Find matching row from LLM output
                matching_row = None
                for r in rows:
                    if int(r.get("score", -1)) == exp_score:
                        matching_row = r
                        break
                        
                if matching_row and matching_row.get("criteria") and matching_row.get("sample"):
                    valid_rows.append(matching_row)
                else:
                    # Fallback criteria row
                    lbl = {
                        3: "Full marks",
                        2: "Full marks" if max_score == 2 else "Partial",
                        1: "Partial marks" if max_score == 2 else "Minimal",
                        0: "Zero"
                    }.get(exp_score, "Criteria tier")
                    
                    valid_rows.append({
                        "score": exp_score,
                        "label": lbl,
                        "criteria": "[RUBRIC ROW INCOMPLETE — PLEASE REVIEW]",
                        "sample": "[Sample response required]"
                    })
                    override_msg = (override_msg + f" WARNING: Rubric row for score {exp_score} missing or invalid. Created fallback.").strip()
                    
            rubric_data["rows"] = valid_rows
            
        # Compile item dictionary
        enriched_item = {
            "Item_ID": item_id,
            "images": row.get("_images", []),
            "Mastery_Level": confirmed_mastery,
            "Blooms_Level": blooms,
            "DoK_Level": dok,
            "Estimated_Time": est_time,
            "Max_Score": max_score,
            "Scoring_Type": scoring_type,
            "If_Partial_Scale": scale,
            "Prerequisite_Level": prereq,
            "Item_Type": item_type,
            "No_Of_Options": num_options,
            "Has_Image": has_image,
            "Image_File_Name": image_file_name,
            "Has_Table": has_table,
            "Has_Equation": has_equation,
            "Equation_Format": eq_format if has_equation == "Yes" else "",
            "Feedback_Type": feedback_type,
            "Topic": item_topic,
            "Topic_ID": item_topic_id,
            "Subject_ID": subject_id,
            "NCF_CG": item_ncf_cg,
            "Competency": item_competency,
            "Learning_Outcome": item_learning_outcome,
            "LO_ID": item_lo_id,
            "Stimulus_Text": stimulus,
            "Item_Stem": stem,
            "Explanation": explanation_text,
            "Rubric": rubric_data,
            # Original excel info
            "Option_A": opt_a,
            "Option_B": opt_b,
            "Option_C": opt_c,
            "Option_D": opt_d,
            "Correct_Option": correct_opt,
            "Correct_Answer": correct_ans,
            "Distractor_Rationales": distractor_rationales,
            "Row_Number": excel_row_num
        }
        
        processed_items.append(enriched_item)
        
        # Log successful audit entry
        status_label = "WARNING" if "WARNING" in override_msg else ("NOTICE" if override_msg else "SUCCESS")
        audit_logs.append({
            "Row_Number": excel_row_num,
            "Item_ID": item_id,
            "Status": status_label,
            "Message": f"Successfully processed item. {override_msg}".strip(),
            "Mastery_Level": confirmed_mastery,
            "Item_Type": item_type,
            "Overrides": override_msg
        })
        
        pass
        
    # ─── 3. GENERATE DOCX FILE ────────────────────────────────────────────────
    if progress_callback:
        progress_callback(total_valid, total_valid, "Finalizing...", "N/A", "Compiling Microsoft Word document (.docx)...")
        
    doc = Document()
    
    # Section margins: 1.5875 cm = 0.625 inches
    section = doc.sections[0]
    section.top_margin = Inches(0.625)
    section.bottom_margin = Inches(0.625)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    
    # Configure Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    
    today_str = datetime.today().strftime("%d/%m/%Y")
    
    for item_idx, item in enumerate(processed_items):
        # Table 1: PART A Header
        t1 = create_styled_table(doc, 1, [9360], border_color="999999")
        set_cell_text(t1.rows[0].cells[0], "PART A — ITEM METADATA  (Tech: used for bulk import and adaptive routing)", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Subsection Header
        add_section_header(doc, "A1.  Item Identity")
        
        # Table 2: Identity Metadata
        t2 = create_styled_table(doc, 3, [1600, 3080, 1600, 3080])
        labels_t2 = [
            [("Item ID *", True), (item["Item_ID"], False), ("Version", True), ("v1.0", False)],
            [("Status *", True), ("Draft", False), ("Date", True), (today_str, False)],
            [("Authored by *", True), ("TechCurators", False), ("Reviewed by", True), ("", False)]
        ]
        for r_idx, row_labels in enumerate(labels_t2):
            for c_idx, (text, is_label) in enumerate(row_labels):
                cell = t2.rows[r_idx].cells[c_idx]
                if is_label:
                    set_cell_shading(cell, "e2e8f0")
                elif r_idx == 2 and c_idx == 3: # Reviewed by empty value cell
                    set_cell_shading(cell, "f4f6f9")
                set_cell_text(cell, text, bold=is_label)
                
        # Subsection Header
        add_section_header(doc, "A2.  Curriculum Alignment")
        
        # Table 3: Alignment Metadata
        t3 = create_styled_table(doc, 5, [1600, 3080, 1600, 3080])
        labels_t3 = [
            [("Grade *", True), ("6", False), ("Subject ID *", True), (item["Subject_ID"], False)],
            [("Chapter *", True), (batch_meta["chapter_name"], False), ("Chapter Code *", True), (batch_meta["chapter_code"], False)],
            [("Topic *", True), (item["Topic"], False), ("Topic ID *", True), (item["Topic_ID"], False)],
            [("NCF CG #", True), (item["NCF_CG"], False), ("Competency", True), (item["Competency"], False)],
            [("Learning Outcome *", True), (item["Learning_Outcome"], False), ("LO ID", True), ("", False)]
        ]
        for r_idx, row_labels in enumerate(labels_t3):
            for c_idx, (text, is_label) in enumerate(row_labels):
                cell = t3.rows[r_idx].cells[c_idx]
                if is_label:
                    set_cell_shading(cell, "e2e8f0")
                set_cell_text(cell, text, bold=is_label)
                
        # Subsection Header
        add_section_header(doc, "A3.  Adaptive Routing  ", "(Tech: drives mastery progression logic)")
        
        # Table 4: Adaptive Routing
        t4 = create_styled_table(doc, 4, [1600, 3080, 1600, 3080])
        labels_t4 = [
            [("Mastery Level *", True), (item["Mastery_Level"], False), ("Bloom's Level *", True), (item["Blooms_Level"], False)],
            [("DoK Level *", True), (item["DoK_Level"], False), ("Estimated Time *", True), (item["Estimated_Time"], False)],
            [("Max Score *", True), (str(item["Max_Score"]), False), ("Scoring Type *", True), (item["Scoring_Type"], False)],
            [("If Partial — Scale", True), (item["If_Partial_Scale"], False), ("Prerequisite Level", True), (item["Prerequisite_Level"], False)]
        ]
        for r_idx, row_labels in enumerate(labels_t4):
            for c_idx, (text, is_label) in enumerate(row_labels):
                cell = t4.rows[r_idx].cells[c_idx]
                if is_label:
                    set_cell_shading(cell, "e2e8f0")
                set_cell_text(cell, text, bold=is_label)
                
        # Subsection Header
        add_section_header(doc, "A4.  Item Format & Rendering  ", "(Tech: platform display and scoring engine)")
        
        # Table 5: Item Format
        t5 = create_styled_table(doc, 4, [1600, 3080, 1600, 3080])
        # Format Item Type name for table display (Rule 3.12 mapping)
        display_type = "Extended Constructed Response" if item["Item_Type"] == "Extended Response" else item["Item_Type"]
        labels_t5 = [
            [("Item Type *", True), (display_type, False), ("No. of Options", True), (item["No_Of_Options"], False)],
            [("Has Image? *", True), (item["Has_Image"], False), ("Image File Name", True), (item["Image_File_Name"], False)],
            [("Has Table? *", True), (item["Has_Table"], False), ("Has Equation? *", True), (item["Has_Equation"], False)],
            [("Equation Format", True), (item["Equation_Format"], False), ("Feedback Type *", True), (item["Feedback_Type"], False)]
        ]
        for r_idx, row_labels in enumerate(labels_t5):
            for c_idx, (text, is_label) in enumerate(row_labels):
                cell = t5.rows[r_idx].cells[c_idx]
                if is_label:
                    set_cell_shading(cell, "e2e8f0")
                set_cell_text(cell, text, bold=is_label)
                
        # Subsection Header
        add_section_header(doc, "A5.  Quality & Review Flags")
        
        # Table 6: Quality Flags
        t6 = create_styled_table(doc, 2, [1600, 3080, 1600, 3080])
        labels_t6 = [
            [("Reviewed for bias?", True), ("Pending", False), ("Reviewed for accuracy?", True), ("Pending", False)],
            [("NCF scope verified?", True), ("Pending", False), ("Cultural check?", True), ("Pending", False)]
        ]
        for r_idx, row_labels in enumerate(labels_t6):
            for c_idx, (text, is_label) in enumerate(row_labels):
                cell = t6.rows[r_idx].cells[c_idx]
                if is_label:
                    set_cell_shading(cell, "e2e8f0")
                set_cell_text(cell, text, bold=is_label)
                
        doc.add_paragraph()
        
        # Table 7: Quality flags guidance
        t7 = create_styled_table(doc, 1, [9360], border_color="2e5f8a")
        set_cell_shading(t7.rows[0].cells[0], "f4f6f9")
        set_cell_text(t7.rows[0].cells[0], "All four quality flags must show 'Yes' before an item can be marked Approved.", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Spacing element
        doc.add_paragraph().paragraph_format.space_before = Pt(4)
        
        # Table 8: PART B Header
        t8 = create_styled_table(doc, 1, [9360], border_color="999999")
        set_cell_text(t8.rows[0].cells[0], "PART B — ITEM CONTENT", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Subsection Header
        add_section_header(doc, "B1.  Stimulus  ", "(the scenario, image, data, or text the learner reads before the question)")
        
        # Table 9: Stimulus guidance
        t9 = create_styled_table(doc, 1, [9360], border_color="2e5f8a")
        set_cell_shading(t9.rows[0].cells[0], "f4f6f9")
        set_cell_text(t9.rows[0].cells[0], "Leave blank if no stimulus. For M1 items a stimulus is often not needed. For M3 and M4 a well-designed stimulus is mandatory.", italic=True)
        
        doc.add_paragraph()
        
        # Table 10: Stimulus text
        t10 = create_styled_table(doc, 1, [2000, 7360])
        set_cell_shading(t10.rows[0].cells[0], "e2e8f0")
        set_cell_text(t10.rows[0].cells[0], "Stimulus text", bold=True)
        stim_cell = t10.rows[0].cells[1]
        set_cell_text(stim_cell, item["Stimulus_Text"])
        
        # If there are images, attach them to the Stimulus block
        if item.get("images"):
            for img_bytes in item["images"]:
                p = stim_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run()
                try:
                    img_stream = io.BytesIO(img_bytes)
                    run.add_picture(img_stream, width=Inches(4.0))
                except Exception as img_err:
                    p_err = stim_cell.add_paragraph()
                    r_err = p_err.add_run(f"[Error rendering attached image: {img_err}]")
                    r_err.font.italic = True
                    r_err.font.color.rgb = RGBColor(255, 0, 0)
        
        # Subsection Header
        add_section_header(doc, "B2.  Item Stem  ", "(the question or instruction the learner must respond to)")
        
        # Table 11: Stem text
        t11 = create_styled_table(doc, 1, [2000, 7360])
        set_cell_shading(t11.rows[0].cells[0], "e2e8f0")
        set_cell_text(t11.rows[0].cells[0], "Item stem *", bold=True)
        set_cell_text(t11.rows[0].cells[1], item["Item_Stem"])
        
        # Subsection Header
        add_section_header(doc, "B3.  Response Options  ", "(for MCQ only — skip for Short Answer / Extended Response)")
        
        # Table 12: Options (Borders 999999 on all cells)
        t12 = create_styled_table(doc, 5, [877, 1059, 7424], border_color="999999")
        
        # Row 1 headers
        set_cell_text(t12.rows[0].cells[0], "Option", bold=True)
        set_cell_text(t12.rows[0].cells[1], "Correct?", bold=True)
        set_cell_text(t12.rows[0].cells[2], "Option Text", bold=True)
        
        if item["Item_Type"] == "MCQ":
            for r_idx, letter in enumerate(["A", "B", "C", "D"], start=1):
                is_correct = "Yes" if item["Correct_Option"] == letter else "No"
                opt_val = item[f"Option_{letter}"]

                # Show diagnostic rationale for ALL options:
                # - Correct option: shows "CORRECT ANSWER: ..." explanation
                # - Incorrect options: shows error-type rationale
                rat = item["Distractor_Rationales"].get(f"rationale_{letter.lower()}", "")
                combined_opt_text = f"{opt_val}\n{rat}" if rat else opt_val

                set_cell_text(t12.rows[r_idx].cells[0], letter, bold=True)
                set_cell_text(t12.rows[r_idx].cells[1], is_correct)
                set_cell_text(t12.rows[r_idx].cells[2], combined_opt_text)
        else:
            # Render empty MCQ table for non-MCQ
            for r_idx, letter in enumerate(["A", "B", "C", "D"], start=1):
                set_cell_text(t12.rows[r_idx].cells[0], letter, bold=True)
                set_cell_text(t12.rows[r_idx].cells[1], "")
                set_cell_text(t12.rows[r_idx].cells[2], "")
                
        # Subsection Header
        add_section_header(doc, "B4.  Correct Answer / Expected Response *")
        
        # Table 13: Correct answer
        t13 = create_styled_table(doc, 1, [2000, 7360])
        set_cell_shading(t13.rows[0].cells[0], "e2e8f0")
        set_cell_text(t13.rows[0].cells[0], "Correct answer", bold=True)
        
        if item["Item_Type"] == "MCQ":
            correct_val = item[f"Option_{item['Correct_Option']}"]
            ans_text = f"{item['Correct_Option']}. {correct_val}"
        else:
            ans_text = item["Correct_Answer"]
            
        set_cell_text(t13.rows[0].cells[1], ans_text)
        
        # Subsection Header
        add_section_header(doc, "B5.  Answer Explanation  ", "(shown to learner as feedback on the platform)")
        
        # Table 14: Explanation
        t14 = create_styled_table(doc, 1, [2000, 7360])
        set_cell_shading(t14.rows[0].cells[0], "e2e8f0")
        set_cell_text(t14.rows[0].cells[0], "Explanation *", bold=True)
        set_cell_text(t14.rows[0].cells[1], item["Explanation"])
        
        # Subsection Header
        add_section_header(doc, "B6.  Marking Rubric  ", "(for Short Answer and Extended Response only)")
        
        # Table 15: Rubric guidance
        t15 = create_styled_table(doc, 1, [9360], border_color="2e5f8a")
        set_cell_shading(t15.rows[0].cells[0], "f4f6f9")
        set_cell_text(t15.rows[0].cells[0], "Skip B6 for MCQ items. For Short and Extended Answer: partial scoring is required — binary marking (full or zero) is not acceptable.", italic=True)
        
        doc.add_paragraph()
        
        # Table 16: Rubric Table (Borders 999999 on all cells)
        rubric_rows_data = item["Rubric"].get("rows", [])
        num_rubric_rows = len(rubric_rows_data) + 1 # +1 for header
        
        if item["Item_Type"] == "MCQ" or num_rubric_rows == 1:
            # MCQ requires 3 blank rows under the header row (4 rows total)
            t16 = create_styled_table(doc, 4, [838, 1562, 6960], border_color="999999")
            set_cell_text(t16.rows[0].cells[0], "Score", bold=True)
            set_cell_text(t16.rows[0].cells[1], "Label", bold=True)
            set_cell_text(t16.rows[0].cells[2], "Criteria — what must be present for this score + sample response", bold=True)
            for r_idx in range(1, 4):
                set_cell_text(t16.rows[r_idx].cells[0], "")
                set_cell_text(t16.rows[r_idx].cells[1], "")
                set_cell_text(t16.rows[r_idx].cells[2], "")
        else:
            t16 = create_styled_table(doc, num_rubric_rows, [838, 1562, 6960], border_color="999999")
            set_cell_text(t16.rows[0].cells[0], "Score", bold=True)
            set_cell_text(t16.rows[0].cells[1], "Label", bold=True)
            set_cell_text(t16.rows[0].cells[2], "Criteria — what must be present for this score + sample response", bold=True)
            for r_idx, r_val in enumerate(rubric_rows_data, start=1):
                criteria_combined = f"Criteria: {r_val['criteria']}\nSample Response: {r_val['sample']}"
                set_cell_text(t16.rows[r_idx].cells[0], str(r_val["score"]))
                set_cell_text(t16.rows[r_idx].cells[1], r_val["label"])
                set_cell_text(t16.rows[r_idx].cells[2], criteria_combined)
                
        # Subsection Header
        add_section_header(doc, "B7.  Reviewer Notes  ", "(internal — not shown to learner)")
        
        # Table 17: Notes
        t17 = create_styled_table(doc, 1, [2000, 7360])
        set_cell_shading(t17.rows[0].cells[0], "e2e8f0")
        set_cell_text(t17.rows[0].cells[0], "Notes", bold=True)
        set_cell_text(t17.rows[0].cells[1], "") # Always empty cell on generation
        
        doc.add_paragraph()
        
        # Table 18: END OF ITEM
        t18 = create_styled_table(doc, 1, [10298], border_color="2e5f8a")
        set_cell_shading(t18.rows[0].cells[0], "f4f6f9")
        set_cell_text(t18.rows[0].cells[0], "END OF ITEM  ·  Duplicate this page for the next item  ·  Do not modify template structure", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Page break after item (except the last one)
        if item_idx < len(processed_items) - 1:
            p_break = doc.add_paragraph()
            p_break.paragraph_format.space_before = Pt(0)
            p_break.paragraph_format.space_after = Pt(0)
            r_break = p_break.add_run()
            r_break.add_break(WD_BREAK.PAGE)
            
    # Save DOCX to Bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    docx_bytes = doc_io.getvalue()
    
    # ─── 4. GENERATE CSV AUDIT LOG ─────────────────────────────────────────────
    csv_io = io.StringIO()
    csv_writer = csv.DictWriter(csv_io, fieldnames=["Row_Number", "Item_ID", "Status", "Message", "Mastery_Level", "Item_Type", "Overrides"])
    csv_writer.writeheader()
    for log in audit_logs:
        csv_writer.writerow(log)
    csv_bytes = csv_io.getvalue().encode("utf-8")
    
    # ─── 5. LOCAL BACKUP SAVE ──────────────────────────────────────────────────
    # Save a local backup inside the workspace 'output' folder
    try:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(excel_file.name)) if hasattr(excel_file, "name") and os.path.exists(excel_file.name) else os.getcwd(), "output")
        os.makedirs(output_dir, exist_ok=True)
        today_file = datetime.today().strftime("%Y%m%d")
        
        backup_docx_path = os.path.join(output_dir, f"Science_PAL_G6_Item_Bank_{today_file}.docx")
        with open(backup_docx_path, "wb") as f:
            f.write(docx_bytes)
            
        backup_csv_path = os.path.join(output_dir, f"Science_PAL_G6_Item_Bank_{today_file}_log.csv")
        with open(backup_csv_path, "wb") as f:
            f.write(csv_bytes)
    except Exception:
        # Ignore local backup failures (e.g. if run purely from memory without filenames)
        pass
        
    summary = {
        "written": len(processed_items),
        "skipped": skipped_items,
        "overrides": override_count
    }
    
    return docx_bytes, csv_bytes, summary